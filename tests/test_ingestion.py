from __future__ import annotations

from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import Image, SimpleDocTemplate

from wald_decision_agent.core.config import AppSettings
from wald_decision_agent.ingestion.ingest import DocumentIngestor


def test_ingestion_supports_docx_and_csv(tmp_path: Path) -> None:
    (tmp_path / "notes.md").write_text("# Update\nRevenue expanded.\n", encoding="utf-8")
    (tmp_path / "metrics.csv").write_text(
        "Department,Performance Score\nSales,88\nSupport,61\n",
        encoding="utf-8",
    )

    doc = Document()
    doc.add_paragraph("Support continued to underperform in the quarter.")
    table = doc.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "Department"
    table.rows[0].cells[1].text = "Margin"
    table.rows[1].cells[0].text = "Sales"
    table.rows[1].cells[1].text = "29"
    table.rows[2].cells[0].text = "Support"
    table.rows[2].cells[1].text = "12"
    doc.save(tmp_path / "board_update.docx")

    ingestor = DocumentIngestor(AppSettings())
    corpus = ingestor.ingest_folder(tmp_path)

    assert any(chunk.source_type == "docx" for chunk in corpus.chunks)
    assert any(chunk.metadata.get("table_id") for chunk in corpus.chunks)
    assert any(table.source_type == "docx_table" for table in corpus.tables.values())


def test_ingestion_uses_ocr_fallback_for_scanned_like_pdf(tmp_path: Path, monkeypatch) -> None:
    image_path = tmp_path / "scan.png"
    image_path.write_bytes(
        b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde\x00\x00\x00\x0cIDATx\x9cc```\x00\x00\x00\x04\x00\x01\xf6\x178U\x00\x00\x00\x00IEND\xaeB`\x82"
    )
    pdf_path = tmp_path / "scan.pdf"
    doc = SimpleDocTemplate(str(pdf_path), pagesize=letter)
    doc.build([Image(str(image_path), width=150, height=150)])

    monkeypatch.setattr(DocumentIngestor, "_extract_pdf_page_with_gemini", lambda self, page: "Scanned revenue memo for Europe")
    corpus = DocumentIngestor(AppSettings()).ingest_folder(tmp_path)

    pdf_chunks = [chunk for chunk in corpus.chunks if chunk.source_path.name == "scan.pdf"]
    assert pdf_chunks
    assert any("Scanned revenue memo for Europe" in chunk.content for chunk in pdf_chunks)
