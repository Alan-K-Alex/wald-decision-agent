from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from openpyxl import Workbook
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = ROOT / "data" / "raw"


def generate_excel() -> None:
    workbook = Workbook()
    actuals = workbook.active
    actuals.title = "Regional Actuals"
    actuals.append(["FY2025 Financial Pack"])
    actuals.append([])
    actuals.append(["Region", "Actual Revenue", "Actual Margin", "Actual Cost"])
    actuals.append(["North America", 102, 31, 52])
    actuals.append(["Europe", 80, 24, 47])
    actuals.append(["APAC", 63, 19, 39])
    actuals.append(["LATAM", 35, 13, 21])

    targets = workbook.create_sheet("Regional Targets")
    targets.append(["Target Plan"])
    targets.append([])
    targets.append(["Region", "Revenue Target", "Margin Target", "Cost Target"])
    targets.append(["North America", 108, 33, 50])
    targets.append(["Europe", 87, 28, 45])
    targets.append(["APAC", 66, 22, 37])
    targets.append(["LATAM", 38, 16, 20])

    risks = workbook.create_sheet("Risk Register")
    risks.append(["Regional Risk Register"])
    risks.append([])
    risks.append(["Region", "Risk Category", "Severity", "Status", "Owner"])
    risks.append(["Europe", "Channel Weakness", "High", "Open", "Sales"])
    risks.append(["APAC", "Partner Coverage", "Medium", "Open", "Partnerships"])
    risks.append(["North America", "Contractor Dependency", "High", "Mitigated", "Operations"])

    workbook.save(DATA_DIR / "board_financial_pack.xlsx")

    messy = Workbook()
    summary = messy.active
    summary.title = "Messy Summary"
    summary.merge_cells("A1:D1")
    summary["A1"] = "Regional KPI Summary"
    summary.append(["Region", "Revenue", None, "Risk"])
    summary.append([None, "Actual", "Target", "Score"])
    summary.append(["North America", 102, 108, 3])
    summary.append(["Europe", 80, 87, 8])
    summary.append(["APAC", 63, 66, 5])
    messy.save(DATA_DIR / "messy_financial_pack.xlsx")


def generate_pdf() -> None:
    doc = SimpleDocTemplate(str(DATA_DIR / "strategy_performance_pack.pdf"), pagesize=letter)
    styles = getSampleStyleSheet()
    story = [
        Paragraph("Strategy Performance Pack", styles["Title"]),
        Spacer(1, 12),
        Paragraph(
            "This pack summarizes regional operating performance. Europe missed both revenue and margin plan, while APAC improved sequentially but remained below target. "
            "North America delivered the largest absolute revenue but still trailed plan. Leadership highlighted open execution risks in Europe and support cost pressure globally.",
            styles["BodyText"],
        ),
        Spacer(1, 12),
        Paragraph("Regional Summary Table", styles["Heading2"]),
        Spacer(1, 8),
    ]
    table = Table(
        [
            ["Region", "Revenue", "Margin", "Plan Status"],
            ["North America", "102", "31", "Below Plan"],
            ["Europe", "80", "24", "Below Plan"],
            ["APAC", "63", "19", "Below Plan"],
            ["LATAM", "35", "13", "Below Plan"],
        ]
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#d0ebff")),
                ("GRID", (0, 0), (-1, -1), 1, colors.black),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ]
        )
    )
    story.extend([table, Spacer(1, 16)])
    story.append(
        Paragraph(
            "Detailed commentary: Europe missed plan because of slower mid-market conversion and weaker channel execution. "
            "APAC improved partner performance but did not close the gap to target. "
            "Leadership recommended preserving strict focus on pricing, partner coverage, and support automation.",
            styles["BodyText"],
        )
    )
    story.append(Spacer(1, 12))
    story.append(
        Paragraph(
            "Additional page content: this paragraph exists to increase document length and to verify multi-page extraction. "
            "The ingestion pipeline should preserve all this content and split it into grounded chunks without losing context.",
            styles["BodyText"],
        )
    )
    story.append(PageBreak())
    story.append(Paragraph("Regional Risk Overview", styles["Heading2"]))
    story.append(Spacer(1, 8))
    risk_table = Table(
        [
            ["Region", "Primary Risk", "Severity", "Owner"],
            ["Europe", "Channel Weakness", "High", "Sales"],
            ["APAC", "Partner Coverage", "Medium", "Partnerships"],
            ["North America", "Support Cost Pressure", "Medium", "Operations"],
        ]
    )
    risk_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#fff3bf")),
                ("GRID", (0, 0), (-1, -1), 1, colors.black),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ]
        )
    )
    story.extend([risk_table, Spacer(1, 16)])
    for _ in range(6):
        story.append(
            Paragraph(
                "Leadership requested that any ingestion system preserve long-form narrative context, table relationships, and supporting commentary without fabricating missing details.",
                styles["BodyText"],
            )
        )
        story.append(Spacer(1, 8))
    doc.build(story)

    fig, ax = plt.subplots(figsize=(7, 3))
    ax.axis("off")
    ax.text(0.05, 0.7, "Scanned Revenue Memo", fontsize=18, weight="bold")
    ax.text(0.05, 0.42, "Europe revenue actual: 80", fontsize=14)
    ax.text(0.05, 0.2, "Europe revenue target: 87", fontsize=14)
    image_path = DATA_DIR / "scanned_revenue_page.png"
    fig.savefig(image_path, dpi=180, bbox_inches="tight")
    plt.close(fig)

    scanned_doc = SimpleDocTemplate(str(DATA_DIR / "scanned_revenue_pack.pdf"), pagesize=letter)
    scanned_story = [Image(str(image_path), width=500, height=220)]
    scanned_doc.build(scanned_story)


def generate_docx() -> None:
    doc = Document()
    doc.add_heading("Operational Steering Memo", level=1)
    doc.add_paragraph(
        "This memo captures departmental execution quality. Support remains the most pressured function, while engineering and sales improved quarter-over-quarter."
    )
    table = doc.add_table(rows=4, cols=3)
    headers = ["Department", "Execution Score", "Commentary"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    rows = [
        ("Sales", "88", "Pipeline discipline improved."),
        ("Engineering", "84", "Release predictability improved."),
        ("Support", "61", "Ticket backlog and contractor dependency remain issues."),
    ]
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            table.rows[row_idx].cells[col_idx].text = value
    doc.save(DATA_DIR / "operational_steering_memo.docx")


def main() -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    generate_excel()
    generate_pdf()
    generate_docx()


if __name__ == "__main__":
    main()
