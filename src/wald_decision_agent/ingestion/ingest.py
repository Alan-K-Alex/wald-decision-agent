from __future__ import annotations

from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from pypdf import PdfReader

from ..core.config import AppSettings
from ..core.logging import get_logger
from ..core.models import Corpus, DocumentChunk, ExtractedDocument, StructuredTable, VisualArtifact
from .pdf_table_extractor import PDFTableExtractor
from .preprocess import split_text_with_offsets
from .spreadsheet_parser import SpreadsheetParser
from ..utils import compact_whitespace, sha256_file, sha256_text
from .visual_extractor import VisualExtractor


class DocumentIngestor:
    def __init__(self, settings: AppSettings) -> None:
        self.settings = settings
        self.logger = get_logger("ingestion.ingest")
        self.spreadsheet_parser = SpreadsheetParser()
        self.pdf_table_extractor = PDFTableExtractor()
        self.visual_extractor = VisualExtractor(settings)

    def ingest_folder(self, folder: str | Path) -> Corpus:
        root = Path(folder)
        corpus = Corpus()
        for path in sorted(root.rglob("*")):
            if not path.is_file():
                continue
            suffix = path.suffix.lower()
            self.logger.debug("Processing file: %s", path)
            if suffix in {".txt", ".md"}:
                document, chunks = self._chunk_text_file(path)
                corpus.documents[document.document_id] = document
                corpus.chunks.extend(chunks)
            elif suffix == ".pdf":
                document, chunks = self._chunk_pdf(path)
                corpus.documents[document.document_id] = document
                corpus.chunks.extend(chunks)
                for table in self.pdf_table_extractor.parse_file(path):
                    corpus.tables[table.table_id] = table
                    corpus.chunks.append(self._table_summary_chunk(table))
            elif suffix == ".docx":
                document, docx_chunks, docx_tables = self._ingest_docx(path)
                corpus.documents[document.document_id] = document
                corpus.chunks.extend(docx_chunks)
                for table in docx_tables:
                    corpus.tables[table.table_id] = table
                    corpus.chunks.append(self._table_summary_chunk(table))
            elif suffix in {".csv", ".tsv", ".xlsx", ".xls"}:
                tables = self.spreadsheet_parser.parse_file(path)
                for table in tables:
                    corpus.tables[table.table_id] = table
                    corpus.chunks.append(self._table_summary_chunk(table))
            elif suffix in {".png", ".jpg", ".jpeg", ".webp", ".svg"}:
                visual = self.visual_extractor.parse_file(path)
                if visual is not None:
                    corpus.visuals[visual.artifact_id] = visual
                    corpus.chunks.append(self._visual_summary_chunk(visual))
        self.logger.info("Folder ingestion complete for %s", root)
        return corpus

    def _chunk_text_file(self, path: Path) -> tuple[ExtractedDocument, list[DocumentChunk]]:
        text = path.read_text(encoding="utf-8")
        return self._chunk_text(text, path=path, source_type="text")

    def _ingest_docx(self, path: Path) -> tuple[ExtractedDocument, list[DocumentChunk], list[StructuredTable]]:
        doc = Document(path)
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
        document, chunks = self._chunk_text(text, path=path, source_type="docx")
        tables: list[StructuredTable] = []
        for index, table in enumerate(doc.tables, start=1):
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if len(rows) < 2:
                continue
            headers = [header if header else f"column_{idx + 1}" for idx, header in enumerate(rows[0])]
            frame = pd.DataFrame(rows[1:], columns=headers).dropna(how="all").reset_index(drop=True)
            frame.insert(0, "_source_row", list(range(2, len(frame) + 2)))
            structured_table = StructuredTable(
                table_id=f"{path.stem}:table:{index}",
                source_path=path,
                source_type="docx_table",
                dataframe=frame,
                metadata={
                    "sheet_name": f"Table {index}",
                    "source_file": path.name,
                    "table_name": f"Table {index}",
                    "columns": [column for column in frame.columns if not str(column).startswith("_")],
                    "source_range": f"Table {index} rows 2-{len(frame) + 1}",
                    "header_rows": 1,
                    "row_count": len(frame),
                },
                retrieval_text=self.spreadsheet_parser._build_retrieval_text(path, f"Table {index}", frame),
            )
            tables.append(structured_table)
        return document, chunks, tables

    def _chunk_pdf(self, path: Path) -> tuple[ExtractedDocument, list[DocumentChunk]]:
        reader = PdfReader(str(path))
        full_pages: list[str] = []
        chunks: list[DocumentChunk] = []
        for page_idx, page in enumerate(reader.pages, start=1):
            page_text = compact_whitespace(page.extract_text() or "")
            if not page_text:
                page_text = self._extract_pdf_page_with_gemini(page)
            full_pages.append(page_text)
            if not page_text:
                continue
            for idx, (content, start, end) in enumerate(self._split_chunks(page_text), start=1):
                chunks.append(
                    DocumentChunk(
                        chunk_id=f"{path.stem}:p{page_idx}:{idx}",
                        source_path=path,
                        content=content,
                        source_type="pdf",
                        metadata={"page": page_idx, "start_offset": start, "end_offset": end, "chunk_index": idx},
                    )
                )
        raw_text = "\n".join(full_pages)
        document = ExtractedDocument(
            document_id=f"{path.stem}:document",
            source_path=path,
            source_type="pdf",
            raw_text=raw_text,
            metadata={"file_hash": sha256_file(path), "text_hash": sha256_text(raw_text), "page_count": len(reader.pages)},
        )
        return document, chunks

    def _extract_pdf_page_with_gemini(self, page) -> str:
        if not self.settings.gemini_api_key:
            return ""
        try:
            from google import genai
            from google.genai import types
        except ImportError:
            return ""

        image_parts: list[object] = []
        for image in getattr(page, "images", []) or []:
            data = getattr(image, "data", None)
            if not data:
                continue
            try:
                image_parts.append(types.Part.from_bytes(data=data, mime_type="image/png"))
            except Exception:
                continue
        if not image_parts:
            return ""
        prompt = (
            "Extract readable business text from this scanned document page. "
            "Return plain text only. Preserve numbers and headings. Do not infer missing values."
        )
        try:
            client = genai.Client(api_key=self.settings.gemini_api_key)
            response = client.models.generate_content(
                model=self.settings.vision_model,
                contents=[prompt, *image_parts],
            )
        except Exception:
            return ""
        text = compact_whitespace(getattr(response, "text", "") or "")
        if text:
            self.logger.info("Used Gemini vision fallback for image-based PDF page")
        return text

    def _chunk_text(self, text: str, path: Path, source_type: str) -> tuple[ExtractedDocument, list[DocumentChunk]]:
        cleaned = compact_whitespace(text)
        document = ExtractedDocument(
            document_id=f"{path.stem}:document",
            source_path=path,
            source_type=source_type,
            raw_text=cleaned,
            metadata={"file_hash": sha256_file(path), "text_hash": sha256_text(cleaned)},
        )
        chunks: list[DocumentChunk] = []
        for idx, (content, start, end) in enumerate(self._split_chunks(cleaned), start=1):
            chunks.append(
                DocumentChunk(
                    chunk_id=f"{path.stem}:{idx}",
                    source_path=path,
                    content=content,
                    source_type=source_type,
                    metadata={"start_offset": start, "end_offset": end, "chunk_index": idx},
                )
            )
        return document, chunks

    def _split_chunks(self, text: str) -> Iterable[tuple[str, int, int]]:
        return split_text_with_offsets(text, self.settings.chunk_size, self.settings.chunk_overlap)

    def _table_summary_chunk(self, table: StructuredTable) -> DocumentChunk:
        return DocumentChunk(
            chunk_id=f"{table.table_id}:summary",
            source_path=table.source_path,
            content=table.retrieval_text,
            source_type="spreadsheet",
            metadata={"table_id": table.table_id, "sheet_name": table.metadata.get("sheet_name")},
        )

    def _visual_summary_chunk(self, visual: VisualArtifact) -> DocumentChunk:
        return DocumentChunk(
            chunk_id=f"{visual.artifact_id}:summary",
            source_path=visual.source_path,
            content=f"{visual.summary} Extracted text: {visual.extracted_text}",
            source_type="visual",
            metadata={"visual_id": visual.artifact_id, "title": visual.metadata.get("title")},
        )
